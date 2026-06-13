"""Extract plain text from various document formats.

Supported formats: PDF, DOCX, TXT, MD, RTF.

PDF extraction reuses the OCR-fallback logic in
core.document_brain.extractor for scanned pages.
"""

from __future__ import annotations

from pathlib import Path

import docx
from striprtf.striprtf import rtf_to_text

from core.document_brain.extractor import extract_document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".rtf"}


class UnsupportedFileTypeError(ValueError):
    pass


def _extract_pdf(path: Path) -> str:
    return extract_document(str(path)).full_text()


def _extract_docx(path: Path) -> str:
    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _extract_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_rtf(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    return rtf_to_text(raw)


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
    ".rtf": _extract_rtf,
}


def extract_text(path: str | Path) -> str:
    """Extract plain text from a document, dispatching on file extension.

    Raises UnsupportedFileTypeError for extensions outside SUPPORTED_EXTENSIONS.
    """
    path = Path(path)
    extension = path.suffix.lower()

    if extension not in _EXTRACTORS:
        raise UnsupportedFileTypeError(f"Unsupported file type: {extension}")

    return _EXTRACTORS[extension](path)


def discover_documents(folder: str | Path, recursive: bool = True) -> list[Path]:
    """Find all files with a supported extension under `folder`."""
    folder = Path(folder)
    pattern = "**/*" if recursive else "*"

    return sorted(
        p for p in folder.glob(pattern)
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )
